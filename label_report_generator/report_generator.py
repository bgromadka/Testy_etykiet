import logging
import os

import fitz
import requests
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from label_report_generator.epl_renderer import EPLRenderer

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")


class LabelReportGenerator:
    def __init__(self,
                 input_labels_folder: str,
                 output_docx: str,
                 number_of_columns_per_page: int = 3,
                 max_rows_per_page: int = 2
                 ):

        self.input_labels_folder = input_labels_folder
        self.output_docx = output_docx
        self.number_of_columns_per_page = number_of_columns_per_page
        self.max_rows_per_page = max_rows_per_page
        self.column_width = {1: Inches(6), 2: Inches(3), 3: Inches(2)}.get(number_of_columns_per_page, Inches(2.5))
        self.allowed_exts = ['.pdf', '.zpl', '.epl']
        self.doc = Document()
        self.epl_renderer = EPLRenderer()

    def _add_pdf_preview(self, filepath, cell):
        try:
            pdf_doc = fitz.open(filepath)
            for page_number in range(min(len(pdf_doc), 1)):
                pix = pdf_doc[page_number].get_pixmap(dpi=120)
                image_path = f'{filepath}_page{page_number + 1}.png'
                pix.save(image_path)
                cell.add_paragraph().add_run().add_picture(image_path, width=self.column_width - Inches(0.3))
            pdf_doc.close()
        except Exception as e:
            cell.add_paragraph(f'Error PDF: {e}')

    def _add_zpl_preview(self, filepath, cell):
        try:
            with open(filepath, 'r', encoding='utf-8') as file:
                label_code = file.read()

            response = requests.post(
                url='http://api.labelary.com/v1/printers/8dpmm/labels/4x6/0/',
                headers={'Accept': 'image/png'},
                data=label_code.encode('utf-8')
            )

            if response.status_code == 200:
                image_path = filepath + '.png'
                with open(image_path, 'wb') as f:
                    f.write(response.content)
                cell.add_paragraph().add_run().add_picture(image_path, width=self.column_width - Inches(0.3))
            else:
                cell.add_paragraph("The label could not be rendered. Code below:")
                cell.add_paragraph(label_code[:500])
        except Exception as e:
            cell.add_paragraph(f'Error ZPL: {e}')

    def _add_epl_preview(self, filepath, cell):
        image_path = self.epl_renderer.render(filepath)
        if image_path and os.path.exists(image_path):
            cell.add_paragraph().add_run().add_picture(image_path, width=self.column_width - Inches(0.3))
        else:
            cell.add_paragraph("Could not render EPL label.")

    def find_files_in_folder_and_subfolders(self):
        file_paths = []
        for root, _, files in os.walk(self.input_labels_folder):
            for file in files:
                if os.path.splitext(file)[1].lower() in self.allowed_exts:
                    file_paths.append(os.path.join(root, file))
        return file_paths

    def add_hyperlink(self, paragraph, text, url):
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0000FF")
        rPr.append(u)
        rPr.append(color)
        new_run.append(rPr)

        t = OxmlElement("w:t")
        t.text = text
        new_run.append(t)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    def set_column_width(self, cell, width_in_inches):
        if not isinstance(width_in_inches, Inches):
            width_in_inches = Inches(width_in_inches)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        twips = int(width_in_inches.inches * 1440)
        tcW.set(qn('w:w'), str(twips))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)

    def add_file_preview_to_cell(self, filepath, cell):
        ext = os.path.splitext(filepath)[1].lower()

        if ext == '.pdf':
            self._add_pdf_preview(filepath, cell)
        elif ext == '.zpl':
            self._add_zpl_preview(filepath, cell)
        elif ext == '.epl':
            self._add_epl_preview(filepath, cell)
        else:
            cell.add_paragraph(f'Unsupported file type: {ext}')

    def generate(self):
        self.doc.add_heading('Label validation report', 0)

        files = self.find_files_in_folder_and_subfolders()
        total_files = len(files)
        logging.info(f"Number of files to be processed: {total_files}")

        for i in range(0, total_files, self.number_of_columns_per_page):
            if (i // self.number_of_columns_per_page) % self.max_rows_per_page == 0 and i != 0:
                self.doc.add_page_break()

            row_files = files[i:i + self.number_of_columns_per_page]
            table = self.doc.add_table(rows=1, cols=len(row_files))
            table.style = 'Table Grid'
            table.autofit = False
            cells = table.rows[0].cells

            for idx, filepath in enumerate(row_files):
                cell = cells[idx]
                self.set_column_width(cell, self.column_width)

                para_title = cell.add_paragraph()
                run = para_title.add_run(os.path.basename(filepath))
                run.bold = True
                run.font.size = Pt(8)

                para_link = cell.add_paragraph()
                self.add_hyperlink(para_link, 'Open file', filepath)

                logging.info(f"File processing: {os.path.basename(filepath)}")
                self.add_file_preview_to_cell(filepath, cell)
                logging.info(f"Processing completed: {idx + i + 1}/{total_files}")

        self.epl_renderer.close()
        self.doc.save(self.output_docx)
        logging.info("Report saved.")
